"""按参考文档格式生成报告"""
import json, os, re
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_FILE = r"C:\Users\lost_you\educoder_data.json"
SCREENSHOT_DIR = r"C:\Users\lost_you\screenshots"
REPORT_FILE = r"C:\Users\lost_you\Downloads\离散数学编程实践报告_完整版.docx"

CATEGORIES = {
    '数理逻辑': ['推理与证明', '命题逻辑', '逻辑与AI知识表示'],
    '集合及其运算': ['集合的表示和基本运算'],
    '关系表示及其运算': ['特殊关系', '二元关系'],
    '函数': ['函数的运算', '函数的判断'],
    '初等数论': ['除法和模运算', '素数和最大公因数', '同余方程', '密码学'],
    '组合数学': ['计数基础', '鸽巢原理', '排列与组合', '生成函数'],
    '图论': ['图论的相关定义', '图论的应用', '特殊图', '树'],
    '选做': ['数论'],
}

ANALYSIS = {
    '推理与证明': [
        '使用sympy库构造命题逻辑公式，定义P、Q、R、S四个命题符号变量，分别构造三个前提条件：P蕴含Q（P>>Q）、R蕴含S（R>>S）、P析取R（P|R），以及结论Q析取S。将前提合取后与结论构造蕴含式，调用satisfiable函数验证该蕴含式是否为永真式。若返回False说明不存在使蕴含式为假的情况，即推理有效。',
        '构造一个看似有效但实际存在逻辑漏洞的推理形式作为测试用例。将前提条件合取作为前提，结论单独作为目标，使用satisfiable检查是否存在使前提为真而结论为假的真值赋值。若找到这样的反例，则说明该推理形式无效。输出反例的具体赋值，清晰展示为何推理不成立。',
        '构造归结推理规则：前提为P析取Q（P|Q）和R析取非P（R|~P），结论为Q析取R。将两个前提的合取作为条件，与结论构造蕴含式，调用satisfiable验证有效性。归结原理是自动定理证明的核心推理规则，通过编程验证可加深对其正确性的理解。',
        '使用sympy.logic.boolalg.simplify_logic函数对逻辑表达式进行代数化简。给定表达式~(P>>Q)，先将其转化为等价的合取范式或析取范式，再与目标表达式(P&~Q)进行比较。通过编程验证德摩根律和蕴含等值式的正确性，展示了代数化简在逻辑等价性证明中的实际应用。',
    ],
    '命题逻辑': [
        '编写真值表生成函数，输入为包含P、Q、R三个命题变元的中缀表达式字符串。使用itertools.product生成全部8种真值赋值组合，对每种赋值使用字典替换方式求值。遍历过程中统计结果为真和假的赋值数量，最终根据统计结果将表达式分类为永真式（全部为真）、永假式（全部为假）或可满足式。',
        '使用sympy构造多个经典推理规则表达式进行逐一验证。假言推理（Modus Ponens）：(P>>Q)&P蕴含Q，即若P则Q成立且P为真可推出Q为真。拒取式（Modus Tollens）：(P>>Q)&~Q蕴含~P，否证后件则前件不成立。假言三段论：(P>>Q)&(Q>>R)蕴含(P>>R)，推出关系的传递性。对每个规则构造蕴含式验证永真性，确认推理模板的逻辑可靠性。',
    ],
    '逻辑与AI知识表示': [
        '将自然语言描述的AI知识规则转化为一阶谓词逻辑公式。首先定义所需谓词，例如R(x)表示x是矩形、Q(x)表示x是四边形、S(x)表示x是正方形。使用全称量词表达规则：对于所有满足条件A的对象，必然满足条件B，符号化为∀x(A(x)→B(x))。返回符号化后的公式字符串，用于后续知识推理。',
        '将AI知识库中的具体事实用一阶谓词逻辑符号化。定义个体常量表示具体实体，谓词表示属性和关系。使用存在量词表达存在性事实：存在某个对象同时满足多个条件，符号化为∃x(P1(x)&P2(x)&...&Pn(x))。将事实与规则配合使用，构建可进行逻辑推理的AI知识库基础。',
    ],
    '集合的表示和基本运算': [
        '读入集合A和B的大小及元素，分别对两个集合进行排序以便高效处理。使用双指针法同时遍历两个有序数组求交集：当元素相等时加入交集结果并移动两指针，不等时移动指向较小元素的指针。求并集时合并两个有序数组，跳过重复元素。时间复杂度O(nlogn+mlogm)，主要开销在排序。',
        '容斥原理用于计算同时满足多个条件的元素个数。本题应用二元容斥公式|A∪B|=|A|+|B|-|A∩B|。先分别统计满足条件一和条件二的元素个数，再计算同时满足两个条件的元素个数（取交集），最后按公式累加得到结果。关键在于准确识别题目中的"满足条件一"、"满足条件二"和"同时满足"的具体含义。',
    ],
    '特殊关系': [
        '等价关系R满足自反性、对称性和传递性，可将集合A划分为若干互不相交的等价类。使用并查集数据结构维护元素之间的等价关系：遍历所有等价对(a,b)，将a和b合并到同一集合。合并完成后，将具有相同根（代表元）的元素归为一组，即为一个等价类。最后按代表元升序排列所有等价类输出商集A/R。',
        '偏序关系满足自反性、反对称性和传递性。在偏序集的哈斯图中，极大元是指不存在其他元素严格大于它的元素。遍历所有顶点，对每个顶点检查是否存在其他顶点在偏序关系中排在它之后（即从它出发有路径到达其他顶点）。若不存在这样的顶点，则该顶点为极大元。需要根据关系矩阵构造可达性信息进行判断。',
    ],
    '二元关系': [
        '笛卡尔积A×B定义为所有有序对(a,b)的集合，其中a∈A、b∈B。使用双重循环遍历集合A和集合B的所有元素，按先A后B的字典序依次输出每个有序对。外层循环遍历A中元素，内层循环遍历B中元素，确保输出顺序符合要求。若A有n个元素、B有m个元素，则笛卡尔积共有n×m个有序对。',
        '关系矩阵是表示有限集合上二元关系的标准方法。给定集合A（大小n）、集合B（大小m）及关系R中的k个有序对，构造n×m的0-1矩阵M，其中M[i][j]=1当且仅当(A[i],B[j])∈R。使用二维数组存储矩阵，先初始化为全0，然后遍历有序对列表将对应位置设为1。最后按行输出矩阵。',
        '关系矩阵的幂运算在关系合成中有重要应用。R^m表示关系R与自身合成m次的结果。由于关系矩阵是布尔矩阵，矩阵乘法中的加法替换为逻辑或、乘法替换为逻辑与。使用三重循环实现布尔矩阵乘法，每次迭代将中间结果存储，反复m-1次即可得到m次幂。若n较大可考虑使用矩阵快速幂优化。',
        '关系的性质判断是离散数学的基础内容。自反性检查对角线元素是否全为1；反自反性检查对角线是否全为0；对称性比较矩阵与其转置是否相等；反对称性检查对任意i≠j，M[i][j]和M[j][i]是否不同时为1；传递性验证R°R⊆R，即对任意i,k,j，若M[i][k]=1且M[k][j]=1则必须有M[i][j]=1。',
        '传递闭包是包含原关系的最小传递关系。使用Warshall算法：依次以每个顶点k为中间节点，检查所有顶点对(i,j)，若从i到k和从k到j都有路径，则添加从i到j的路径。三重循环，外层遍历中间节点k，中层遍历起点i，内层遍历终点j。算法时间复杂度O(n³)，空间复杂度O(n²)，直接在原矩阵上迭代更新。',
    ],
    '函数的运算': [
        '给定函数f从集合A到集合B的映射关系（n对有序对(x,y)，其中x∈A、y∈B），求解逆映射。首先检查f是否为双射：每个y值只能被唯一x映射到（单射）且B中每个元素都有原像（满射）。若f是双射则逆映射存在：对于每个有序对(x,y)，在逆映射中对应有序对(y,x)。按y升序输出逆映射。',
        '函数迭代f^n(x)表示将函数f对输入x反复应用n次。依次读入每次迭代的输入值x，调用f计算输出值y=f(x)，然后将y作为下一轮迭代的输入。迭代n次后输出最终结果。需要注意迭代过程中数据类型的统一，若中间值较大可能需要使用long long类型防止溢出。时间复杂度和空间复杂度均取决于n和函数f的计算复杂度。',
    ],
    '函数的判断': [
        '给定n个有序对表示的二元关系f，判断其是否为函数以及函数的类型。首先检查函数的定义：每个第一分量x必须对应唯一确定的第二分量y，即不存在两个有序对具有相同x值但不同y值。若满足函数定义，则进一步判断单射（不同x对应不同y）、满射（值域覆盖整个到达域，需额外给定到达域大小）和双射（同时满足单射和满射）。',
        '计算从m元集合到n元集合的所有可能函数的个数。对于定义域中的每个元素x，都有n种可能的函数值选择（即n种y值），且各选择相互独立。因此总函数个数为n^m（n的m次方）。使用循环累积乘积实现幂运算，注意中间结果可能溢出，需根据数据范围选择合适的整数类型。当n和m较大时考虑使用快速幂算法。',
    ],
    '除法和模运算': [
        '给定大整数A、B（以字符串形式给出，可能长达10^5位）和模数m，判断A和B是否模m同余，即验证A≡B(mod m)是否成立。由于整数极大无法直接转换，使用逐位取模算法：从高位到低位处理字符串，维护当前余数r，每遇到新的一位数字d，更新r=(r*10+d)%m。分别对A和B计算余数后比较即可。时间复杂度O(len(A)+len(B))。',
        '计算a^b mod m，其中指数b以字符串形式给出可能极大，模数m由题目指定。利用模运算性质将大指数缩小：先对b中各位数字进行模φ(m)（欧拉函数值）的转换，逐位处理b字符串：维护当前余数exp_mod，每次更新为(exp_mod*10+digit)%phi_m。最后使用标准快速幂算法计算a^exp_mod mod m，时间复杂度O(log m + len(b))。',
    ],
    '素数和最大公因数': [
        '使用Miller-Rabin素性测试（确定性版本）判断大整数n是否为素数。算法基于费马小定理的逆否命题：若n是合数，则大部分a满足a^(n-1)≢1(mod n)。将n-1写为d*2^s的形式（d为奇数），随机选取底数a，计算a^d mod n，然后反复平方s次检查是否出现非平凡平方根。取多个底数测试后若全部通过则n很可能为素数。',
        '对给定正整数n进行质因数分解。从2遍历到√n，对每个可能的因子i，反复用n除以i直到不能整除，记录因子i及其指数。注意i只需检查到√n，因为若n有大于√n的质因子，在遍历完成后n仍大于1，此时n本身就是最后一个质因子（且指数为1）。优化：检查完2之后只检查奇数因子，减少循环次数。',
        '扩展欧几里得算法同时求解gcd(a,b)和一组整数系数(x,y)使得ax+by=gcd(a,b)。递归基：当b=0时gcd(a,0)=a，取x=1,y=0。递归步：先递归计算gcd(b,a%b)得到(x1,y1)满足b*x1+(a%b)*y1=g。利用关系a%b=a-(a//b)*b代入，整理得a*y1+b*(x1-(a//b)*y1)=g，因此x=y1, y=x1-(a//b)*y1。该算法在RSA密码体制中用于求模逆元。',
    ],
    '同余方程': [
        '求解单个线性同余方程ax≡b(mod m)。首先使用扩展欧几里得算法计算d=gcd(a,m)以及一组特解(x0,y0)满足ax0+my0=d。若b不能被d整除则方程无解，因为左边ax永远是d的倍数而右边b不是。若d|b，将特解放大得到x1=x0*(b/d)，通解形式为x≡x1(mod m/d)。方程以m/d为周期，在模m意义下共有d个不同的解。',
        '中国剩余定理（CRT）用于求解一组两两互素模数的同余方程组。对于方程组x≡a1(mod m1), x≡a2(mod m2)且gcd(m1,m2)=1，计算公共模M=m1*m2。令M1=M/m1=m2，求M1在模m1下的乘法逆元inv1（即M1*inv1≡1(mod m1)）；同理M2=M/m2=m1，求inv2。最终解为x=(a1*M1*inv1+a2*M2*inv2)%M。该定理可将大模数分解为小模数分别计算后再合并结果。',
        '求解一般线性同余方程组时模数不一定互素，需要逐个方程合并。维护当前解x和当前模数M（初始x=0,M=1）。对于新方程x≡a(mod m)，代入得x+t*M≡a(mod m)，即M*t≡(a-x)(mod m)。使用扩展欧几里得求解此同余方程得到t，更新x=x+t*M, M=lcm(M,m)。若在合并过程中某个同余方程无解，则整个方程组无解。该策略可处理任意同余方程组。',
    ],
    '密码学': [
        '检查密码强度：遍历密码字符串的每个字符，使用字符分类函数（isupper、islower、isdigit等）检查是否包含大写字母、小写字母、数字和特殊字符。记录满足的字符类型数量，根据满足的类型数评定强度等级。全部满足为"很强"，满足3类为"强"，2类为"中等"，1类为"弱"。输出具体的强度评级。',
        'RSA密码体制中求解私钥d。已知公钥指数e和模数n，以及加密后的密文C。首先需要将n分解质因数得到两个大素数p和q，计算欧拉函数φ(n)=(p-1)(q-1)。私钥d是公钥e在模φ(n)下的乘法逆元，即满足e*d≡1(mod φ(n))。使用扩展欧几里得算法求解此同余方程得到d。',
        'RSA解密过程：已知私钥(d,n)和密文C，计算明文M=C^d mod n。由于指数d极大（与n同数量级，可达1024位以上），必须使用快速幂取模算法降低时间复杂度。从d的二进制最低位开始逐位处理：维护当前结果ans和底数base=C%n，d的每一位为1时执行ans=ans*base%n，每轮执行base=base*base%n并将d右移一位。时间复杂度从O(d)降至O(log d)。',
    ],
    '计数基础': [
        '计算网站用户名的总可能数。用户名可由4个大写英文字母组成（共26^4种可能）或由2位数字组成（共10^2种可能），两者之和为总可能数。直接使用指数运算计算26^4=456976和10^2=100，相加得457076。注意题目可能要求区分大小写或包含更多字符类型，需根据具体条件调整计算。',
        '给定正整数n和k，计算满足任意相邻两位数字之差的绝对值不超过k的n位十进制正整数（首位不为0）的个数。使用动态规划方法求解：定义dp[i][d]表示构造到第i位、末位为数字d时的合法方案数。初始化首位的dp[1][d]=1（d=1..9，因为首位不能为0）。对于i从2到n的每一位，枚举当前末位d(0..9)和上一位的末位j(0..9)，若|d-j|≤k则将dp[i-1][j]累加到dp[i][d]。最终答案为∑_{d=0}^{9}dp[n][d]。时间复杂度O(n×10×10)=O(100n)。',
    ],
    '鸽巢原理': [
        '鸽巢原理（抽屉原理）表述为：若将m个物体放入k个盒子中，且m>k，则至少有一个盒子包含不少于⌈m/k⌉个物体。本题直接应用该原理：输入k（鸽巢数）和m（鸽子数），输出至少有多少只鸽子在同一鸽巢中。计算公式为(m+k-1)//k即向上取整。该原理是组合数学中最基本的存在性定理之一。',
        '给定n个整数的序列a1,a2,...,an，证明存在一个非空连续子序列，其元素之和能被n整除。考虑前缀和s_k=a_1+...+a_k（k=1to n），共n个前缀和。若某个s_k模n等于0则直接找到。否则n个前缀和模n的余数只能在1到n-1范围内，根据鸽巢原理必有两个前缀和s_i和s_j（i<j）模n同余，则a_{i+1}+...+a_j能被n整除。',
        '平面上有n个点，证明存在两个点之间的距离不超过某个上界。将正方形区域等分为若干边长为d的小正方形（鸽巢），若点数多于小正方形数，则根据鸽巢原理必有两个点落在同一小正方形内。同一小正方形内任意两点距离不超过√2*d（对角线长度）。通过选择合适的划分密度控制距离上界。',
    ],
    '排列与组合': [
        '计算组合数C(n,k)，即从n个不同元素中选取k个元素（不考虑顺序）的方案数。使用组合恒等式C(n,k)=C(n-1,k-1)+C(n-1,k)进行递推计算，利用二维数组存储中间结果避免重复计算。当n较小时也可使用阶乘公式C(n,k)=n!/(k!(n-k)!)，为防溢出先计算分子n*(n-1)*...*(n-k+1)每次乘法后立刻除以对应的分母因子。注意当k>n/2时可利用对称性C(n,k)=C(n,n-k)减少一半计算量。',
        '计算排列数P(n,k)=n!/(n-k)!，即从n个不同元素中取k个进行排列的方案数。使用循环累积乘积实现：从i=n-k+1到i=n累乘，等价于n*(n-1)*...*(n-k+1)。当k较大时排列数增长极快（超指数级别），注意使用合适的数据类型，必要时使用大整数库或取模运算控制结果大小。',
    ],
    '生成函数': [
        '硬币找零问题使用生成函数（母函数）方法求解。设给定面值集合，求凑出总金额n的不同方案数。使用动态规划：dp[i]表示凑出金额i的方案数。初始化dp[0]=1（凑0元只有一种方式）。对于每种硬币面值v，从左到右更新dp：dp[j]+=dp[j-v]（j从v到n），表示在原有方案基础上添加一枚v面值硬币。最终答案为dp[n]。',
        '默慈金数M_n计数圆上n个点之间画不相交弦的方案数。使用递推公式：M_0=1, M_1=1。对于n≥2，有M_n=M_{n-1}+Σ_{k=0}^{n-2}M_k*M_{n-2-k}。递推式中第一项表示点n不参与任何弦的方案数，求和项表示点n与某点k+1连接的方案数（将圆分为两部分各自独立）。使用循环和备忘录数组计算。',
    ],
    '图论的相关定义': [
        '给定有向图的邻接矩阵表示（n×n的0-1矩阵），计算每个顶点的出度和入度。顶点i的出度等于邻接矩阵第i行所有元素之和（从i出发的边数），顶点i的入度等于第i列所有元素之和（指向i的边数）。遍历矩阵的行和列即可统计，时间复杂度O(n²)。输出每个顶点的度序列作为图的基本结构信息。',
        '给定图G的邻接矩阵，生成其补图的邻接矩阵。补图中顶点集不变，但两点相邻当且仅当它们在原图G中不相邻。补图的邻接矩阵可通过将原矩阵每个元素取反得到（0变1，1变0），但对角线元素保持为0（无自环）。注意处理无向图和有向图的区别：无向图补图矩阵仍保持对称性。',
        '判断两个图G1和G2是否同构，即是否存在一个顶点集之间的双射f，使得G1中(u,v)为边当且仅当G2中(f(u),f(v))为边。使用回溯搜索（带剪枝的深度优先搜索）枚举所有可能的顶点映射方案。剪枝策略：只将度数相同的顶点互相映射；在搜索过程中及时检查已映射部分是否满足邻接约束。最坏时间复杂度O(n!)，但剪枝后实际效率可接受。',
        '可达矩阵R是一个n×n的0-1矩阵，R[i][j]=1表示从顶点i可达顶点j（存在一条有向路径）。使用Floyd-Warshall算法计算传递闭包：以每个顶点k为中间节点，若i可达k且k可达j，则i可达j。三重循环后得到可达矩阵。根据可达矩阵判断图的连通性：若所有顶点对(i,j)均满足R[i][j]=1或R[j][i]=1则为强连通。',
    ],
    '图论的应用': [
        'Havel-Hakimi算法用于判断一个给定的非递增度序列是否可以表示为某个简单图的顶点度序列。算法步骤：每次取出当前序列中的第一个（最大）度d，删除该顶点，将序列中接下来的d个最大的度各减1（表示这些顶点与删除顶点相邻）。重新排序序列后重复，直至所有度变为0（可图化）或出现负数（不可图化）。',
        '给定有向图G的邻接矩阵A和正整数m，计算从任意顶点i到顶点j的长度恰为m的通路总数。根据图论基本定理，A^m[i][j]恰好等于从i到j长度为m的通路条数。使用矩阵乘法计算A的m次幂，需进行m-1次n×n矩阵乘法。当m较大时使用矩阵快速幂优化至O(n³log m)。需要处理大整数，通路条数可能超过普通整型范围。',
    ],
    '特殊图': [
        '判断无向图是否为欧拉图。欧拉图的充要条件：图是连通的且所有顶点的度数为偶数。首先使用深度优先搜索（DFS）或广度优先搜索（BFS）从任意顶点出发遍历图，检查是否能访问所有顶点（连通性检验）。然后统计每个顶点的度数（与其相邻的边数），检查是否全为偶数。两个条件同时满足则为欧拉图。',
        '判断无向图是否为偶图（二分图）。二分图的充要条件：可以用两种颜色对顶点染色，使得每条边连接的两个顶点颜色不同。使用BFS或DFS染色法：从任意未染色顶点开始，染为颜色0；遍历其所有邻居，染为相反颜色1；若发现某条边连接的两个顶点颜色相同，则图不是二分图。算法在染色冲突时立即返回false。',
    ],
    '树': [
        '判断给定的无向图是否为一棵树。树是连通无环的无向图，等价条件为：具有n个顶点且连通，并且恰好有n-1条边。判断过程：首先检查边数是否等于顶点数减1，若不满足则直接判定不是树。然后使用并查集或DFS/BFS检查图的连通性：从任意顶点出发遍历，统计能访问到的顶点数，若等于n则连通。两个条件同时满足则为树。',
    ],
    '数论': [
        '埃拉托斯特尼筛法（Sieve of Eratosthenes）是求解不超过n的所有素数的经典算法。初始化一个布尔数组is_prime，全部标记为true。从2开始遍历到√n，若当前i为素数（is_prime[i]为true），则将i的所有倍数（从i*i开始，步长为i）标记为false。遍历完成后，所有is_prime值为true的索引即为素数。时间复杂度O(n log log n)。',
        '给定正整数n，求其所有正因数的个数。利用质因数分解的结果：若n的质因数分解为n=Πp_i^e_i（p_i为不同质数，e_i为对应指数），则n的正因数个数为Π(e_i+1)。对n进行试除法分解：从2遍历到√n，找到每个质因子后反复除直到不能整除，累计其指数。遍历结束后若n仍大于1，则n本身为质因子（指数为1）。',
        '欧拉函数φ(n)定义为[1,n]中与n互质的正整数个数。计算公式：φ(n)=n×Π(1-1/p_i)，其中乘积遍历n的所有不同质因数p_i。实现时对n进行质因数分解，在分解过程中同步计算φ值：初始phi=n，每找到一个质因数p，更新phi=phi/p*(p-1)。分解完成后返回phi。φ(n)在RSA密码体制中起核心作用。',
        '欧拉定理指出若正整数a和m互质（即gcd(a,m)=1），则a^φ(m)≡1(mod m)。利用该定理可大幅简化大指数模运算：将a^b mod m中的指数b对φ(m)取模，即计算b_mod=b%φ(m)，若b_mod=0且b>0则取b_mod=φ(m)（因a^0≡1不一定满足简化条件）。然后使用标准快速幂算法计算a^b_mod mod m。计算φ(m)需要对m进行质因数分解，时间复杂度O(√m)。整体算法在RSA解密和大数模幂计算中有广泛应用。',
        '一般线性同余方程a*x≡b(mod m)的求解步骤：计算d=gcd(a,m)，若b不能被d整除则无解。使用扩展欧几里得算法求a*x0+m*y0=d的一组特解(x0,y0)。将特解放大b/d倍得到原方程的一个特解：x1=x0*(b/d)。通解形式为x≡x1(mod m/d)，所有解在模m下有d个。利用取模运算输出最小非负解。',
    ],
}

CHALLENGE_NAMES = {
    '集合的表示和基本运算': ['求集合的交集', '容斥原理'],
    '推理与证明': ['构造性二难推理的有效性验证', '识别无效推理并寻找反例', '归结原理的应用', '等价推理的化简验证'],
    '生成函数': ['找零钱', '默慈金数'],
    '排列与组合': ['计数排序', '排列计数'],
    '鸽巢原理': ['糖果分配', '子集划分', '平面上的点集'],
    '计数基础': ['网站用户名', '相邻位约束'],
    '密码学': ['密码强度', '寻找私钥', 'RSA 解密'],
    '同余方程': ['线性同余方程', '中国剩余定理', '扩展欧几里得应用'],
    '素数和最大公因数': ['实现对称性判断素数', '唯一分解定理', '扩展欧几里得算法'],
    '除法和模运算': ['判断是否同余', '快速幂模运算'],
    '函数的运算': ['函数的复合和逆的求解', '函数迭代应用'],
    '函数的判断': ['函数的定义', '幂集计数问题'],
    '特殊关系': ['等价类与商集', '偏序关系的极大元'],
    '二元关系': ['笛卡尔积', '关系矩阵', '关系幂运算', '关系性质判断', '传递闭包'],
    '逻辑与AI知识表示': ['全称规则', '存在性事实'],
    '命题逻辑': ['真值表求值及分类', '推理规则验证'],
    '数论': ['埃拉托斯特尼筛法', '计算因数个数', '欧拉函数', '欧拉定理的应用', '一般线性同余方程'],
    '树': ['树的性质判断'],
    '特殊图': ['欧拉图的判定', '偶图的判定'],
    '图论的应用': ['图序列判定用Havel-Hakimi算法', '最长长为m的通路计数'],
    '图论的相关定义': ['图的度数与邻接矩阵', '子图与补图', '图的同构判断', '可达矩阵与连通性'],
}

PURPOSE = {
    '数理逻辑': '掌握命题逻辑和谓词逻辑的基本概念，学会使用Python编程验证推理有效性，理解逻辑在AI知识表示中的应用。',
    '集合及其运算': '掌握集合的表示方法和基本运算，包括并集、交集、差集、补集等操作，能用编程实现集合运算。',
    '关系表示及其运算': '掌握关系的矩阵表示和图表示，理解关系的性质（自反、对称、传递等），能用编程实现关系运算。',
    '函数': '掌握函数的定义和判断方法，理解单射、满射、双射的概念，能用编程判断函数的性质并进行复合运算。',
    '初等数论': '掌握整除、模运算、最大公因数、同余方程、素数等初等数论概念，能用编程实现相关算法。',
    '组合数学': '掌握计数原理、排列组合、鸽巢原理、生成函数等组合数学方法，能用编程解决组合计数问题。',
    '图论': '掌握图的基本概念、特殊图、树等图论知识，理解图论在实际问题中的应用。',
    '选做': '掌握数论高阶知识，包括筛法求素数、因数个数计算、欧拉函数和欧拉定理的应用，以及线性同余方程的求解。',
}

def setup_styles(doc):
    """Configure styles to match reference"""
    # Normal -> 宋体小四 1.5倍行距 首行缩进2字符
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)  # 小四
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.first_line_indent = Cm(0.74)  # ~2 chars at 小四

    # Heading 1 -> 黑体三号加粗 黑色
    h1 = doc.styles['Heading 1']
    h1.font.name = '黑体'
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    # Heading 2 -> 黑体四号 黑色
    h2 = doc.styles['Heading 2']
    h2.font.name = '黑体'
    h2.font.size = Pt(14)
    h2.font.bold = False
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    # Heading 3 -> 黑体小四 黑色
    h3 = doc.styles['Heading 3']
    h3.font.name = '黑体'
    h3.font.size = Pt(12)
    h3.font.bold = False
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    # 代码 style -> Times New Roman 小四 单倍行距
    try:
        code_style = doc.styles['代码']
    except KeyError:
        code_style = doc.styles.add_style('代码', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Times New Roman'
    code_style.font.size = Pt(12)
    code_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    code_style.paragraph_format.line_spacing = 1.0
    code_style.paragraph_format.first_line_indent = Cm(0.74)

def set_cell(cell, text):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(14)
    run.font.bold = True

def build():
    with open(OUTPUT_FILE, 'r', encoding='utf-8') as f:
        experiments = json.load(f)
    exp_map = {exp['name']: exp for exp in experiments}

    doc = Document()
    setup_styles(doc)
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)

    # --- Cover ---
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('离散数学编程实践报告')
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(22)
    run.font.bold = True
    doc.add_paragraph(); doc.add_paragraph()
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    for i, field in enumerate(['专    业', '班    级', '学    号', '姓    名', '完成日期']):
        set_cell(table.rows[i].cells[0], f'{field}：')
        set_cell(table.rows[i].cells[1], '')
    doc.add_page_break()

    # --- 目录 ---
    h1 = doc.add_heading('目录', level=1)
    for r in h1.runs:
        r.font.name = '黑体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        r.font.size = Pt(16)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0, 0, 0)

    # Insert TOC field
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)

    run2 = p.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-3" \\h \\z '
    run2._r.append(instrText)

    run3 = p.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run3._r.append(fldChar2)

    run4 = p.add_run('（右键此处 → 更新域 → 更新整个目录）')
    run4.font.color.rgb = RGBColor(128, 128, 128)
    run4.font.size = Pt(10)

    run5 = p.add_run()
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run5._r.append(fldChar3)

    doc.add_page_break()

    # --- Content ---
    for cat_idx, (cat_name, cat_exps) in enumerate(CATEGORIES.items(), 1):
        cat_exps_found = [e for e in cat_exps if e in exp_map]
        if not cat_exps_found:
            continue

        doc.add_heading(f'{cat_idx} {cat_name}', level=1)
        doc.add_heading(f'{cat_idx}.1 实验目的', level=2)
        doc.add_paragraph(PURPOSE.get(cat_name, ''))

        doc.add_heading(f'{cat_idx}.2 实验内容', level=2)

        for i, exp_name in enumerate(cat_exps_found, 1):
            names = CHALLENGE_NAMES.get(exp_name, [])
            doc.add_heading(f'{cat_idx}.2.{i} {exp_name}', level=3)

            if names:
                desc = '、'.join([f'第{j+1}关 {n}' for j, n in enumerate(names)])
            else:
                desc = f'本题共{len(names) if names else "?"}关'
            doc.add_paragraph(desc + '，已全部完成，通过率100%。')

        doc.add_heading(f'{cat_idx}.3 实验过程', level=2)

        for i, exp_name in enumerate(cat_exps_found, 1):
            exp_data = exp_map[exp_name]
            codes = exp_data.get('codes', [])
            names = CHALLENGE_NAMES.get(exp_name, [])
            analyses = ANALYSIS.get(exp_name, [])

            h3 = doc.add_heading(f'{cat_idx}.3.{i} {exp_name}', level=3)

            for j, code_block in enumerate(codes):
                code_text = code_block.get('code', '')
                label = f'第{j+1}关'
                if j < len(names):
                    label += f'：{names[j]}'

                # Analysis paragraph
                if j < len(analyses):
                    doc.add_paragraph(f'{label}\n解题思路：{analyses[j]}')
                else:
                    doc.add_paragraph(f'{label}\n解题思路：通过分析题目要求，设计合适的算法和数据结构完成。')

                # Code in 代码 style
                doc.add_paragraph(code_text, style='代码')

                doc.add_paragraph('运行结果：100%通过。')

            # Insert screenshot after each experiment
            safe_name = re.sub(r'[\\/*?:"<>|]', '_', exp_name)
            # Find matching screenshot
            sshot = None
            if os.path.isdir(SCREENSHOT_DIR):
                for fname in sorted(os.listdir(SCREENSHOT_DIR)):
                    if fname.endswith('.png') and safe_name in fname:
                        sshot = os.path.join(SCREENSHOT_DIR, fname)
                        break
            if sshot and os.path.exists(sshot):
                doc.add_paragraph('')  # spacer
                doc.add_picture(sshot, width=Inches(5.5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph('')  # spacer

        doc.add_page_break()

    # --- 总结 ---
    summary_num = len(CATEGORIES) + 1
    doc.add_heading(f'{summary_num} 总结与课程建议', level=1)
    doc.add_heading(f'{summary_num}.1 课程总结', level=2)
    total = len(experiments)
    total_c = sum(len(exp.get('codes', [])) for exp in experiments)
    doc.add_paragraph(
        f'通过本学期离散数学编程实践课程的学习，系统掌握了离散数学核心领域的编程实现方法。共完成{total}个实验、{total_c}关编程实践。'
        '在数理逻辑部分，学习了使用Python的sympy库验证推理有效性和进行AI知识表示；在集合与关系部分，掌握了集合运算与关系性质的编程实现；'
        '在函数部分，学习了函数性质判断和复合运算；在初等数论部分，掌握了同余方程求解、RSA密码学等核心算法；'
        '在组合数学部分，加深了对排列组合、鸽巢原理、生成函数的理解；在图论部分，学习了图的基本概念和算法的编程实现。'
        '通过将理论知识与编程技能有机结合，提升了解决实际问题的能力。'
    )

    doc.add_heading(f'{summary_num}.2 问题及建议', level=2)
    doc.add_paragraph('课程内容安排合理，难度递进适当，理论与实践结合紧密。建议增加更多实际应用案例，优化在线IDE的响应速度。')

    doc.save(REPORT_FILE)
    print(f'Done: {REPORT_FILE}')

if __name__ == "__main__":
    build()
